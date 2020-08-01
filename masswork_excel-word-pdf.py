#!/usr/bin/env python
# coding: utf-8

# In[1]:


#pip install python-docx


# In[5]:


from docx import Document
import xlrd


# In[6]:


def text_change(headline, data):
    # 用来替换word段落中的关键字内容，关键字都是excel表格的标题行
    myparagraphs = document.paragraphs
    for paragraph in myparagraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(headline, data)#每行数据
            run.text = run_text

    # 用来替换word表格中的关键字内容，关键字都是excel表格的标题行
    mytables = document.tables
    for table in mytables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(headline, data)
                cell.text = cell_text


# In[7]:


#打开excel文件，获取报告内容
xlsx = xlrd.open_workbook(r'F:\LEARN_Python\办公自动化\data.xlsx')
table = xlsx.sheet_by_index(0)#读取了第一个sheet


# In[8]:


#遍历excel的单元格，同时打开报告模板文件，按照excel中的数据替换报告模板中的关键字
#替换完成后，保存为新文件，文件名为excel中的A列单元格的内容+eSRVCC切换成功率低优化报告.docx
for table_row in range(1, table.nrows):#每一行数据读取
    document = Document(r'F:\LEARN_Python\办公自动化\report.docx')#可以改进成一个变量，写入def text_chenge(headline, data,route)里面
    for table_col in range(0, table.ncols):#读取每一列，0是首行
        text_change(str(table.cell(0, table_col).value), str(table.cell(table_row, table_col).value))#excel的data转换成word里的字
        # 将excel表格中的内容替换掉标题行，因为标题行即为报告模板中的关键字
    document.save('F:\\LEARN_Python\\办公自动化\\'+f'{str(table.cell(table_row, 0).value)} 季度分红报告.docx')
    print("%s 季度分红报告生成成功！" % str(table.cell_value(table_row, 0)))


# In[9]:


# doc2pdf.py: python script to convert docx to pdf 
# Requires python for win32 extension
import sys, os
from win32com.client import Dispatch
from os import walk

wdFormatPDF = 17


# In[10]:


#doc --> pdf
def doc2pdf(input_file):
    word = Dispatch('Word.Application')
    doc = word.Documents.Open(input_file)
    doc.SaveAs(input_file.replace(".docx", ".pdf"), FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()


# In[11]:


#遍历整个文件夹，将所有word文件转为pdf；10个/min
if __name__ == "__main__":
    doc_files = []
    directory = "F:\LEARN_Python\办公自动化"
    for root, dirs, filenames in walk(directory):
        for file in filenames:
            if file.endswith(".doc") or file.endswith(".docx"):
                doc2pdf(str(root + "\\" + file))

